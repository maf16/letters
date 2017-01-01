import shutil
import docx
import re
from openpyxl import load_workbook
import string

class Source():                                             #class which contains variable_find function and saves vars to xlsx sheet
    def __init__(self, source_path = 'Vermieterbrief.docx', data_path = 'Data.xlsx'):
        self.handle = docx.Document(source_path)
        self.variables = []
        self.data_path = data_path
    def vars_find(self):                                    #finds variables in xdocx
        for paragraph in self.handle.paragraphs:
            occurances = re.findall("\%\w+", paragraph.text)
            for occurance in occurances:
                self.variables.append(occurance)
        return self.variables
    def vars_write_to_excel(self):
        self.vars_find()
        wb = load_workbook(self.data_path)
        ws = wb['printing']
        counter = 3
        for variable in self.variables:
            string = list(variable)
            del (string[0])
            string = "".join(string)
            ws['A' + str(counter)] = string
            counter += 1
        while counter < 50:
            ws['A' + str(counter)] = ''
            counter += 1
        wb.save(self.data_path)
        return self.variables
class Target():                                     #class which contains replacements_get function and creates final document, using list of vars
    def __init__(self,variables = [],source_path = 'Vermieterbrief.docx',target_path = 'Mieterbriefe.docx',data_path = 'Data.xlsx'):
        shutil.copy2(source_path,target_path)
        self.target_path = target_path
        self.target_handle = docx.Document(target_path)
        self.source_handle = docx.Document(source_path)
        self.variables = variables
        self.replacements = []
        wb = load_workbook(data_path)
        self.ws = wb['printing']
    def empty(self):
        for paragraph in self.target_handle.paragraphs:
            p = paragraph._element                                      #potentially simplify
            p.getparent().remove(p)
            p._p = p._element = None
        self.target_handle.save(self.target_path)
    def read_col(self,col):
        if col > 26:
            raise IndexError('read_col: col > 26')
        counter = 3                                         #start at 3rd row -> layout of xlsx sheet
        tennant_list = []
        while counter < (len(self.variables) + 3):
            text = self.ws[string.ascii_uppercase[col] + str(counter)].value
            if text == None:
                text = ' '
            tennant_list.append(text)
            counter += 1
        return tennant_list
    def read_replacements(self):
        counter = 1                                                     #starting with col(B) in xlsx sheet
        while not self.ws[string.ascii_uppercase[counter] + '3'].value == None:
            text = self.read_col(counter)
            if text == None:
                text = ' '
            self.replacements.append(text)
            counter += 1
        return self.replacements
    def compile(self):
        for tennant in self.replacements:
            for paragraph in self.source_handle.paragraphs:
                counter = 0
                text = paragraph.text
                for variable in self.variables:
                    text = text.replace(variable, tennant[counter])
                    counter += 1
                self.target_handle.add_paragraph(text)
                #Doesnt account for single words in bold or any form of underline
                for run in paragraph.runs:
                    if run.bold:
                        self.target_handle.paragraphs[-1].runs[0].bold = True
        self.target_handle.save(self.target_path)